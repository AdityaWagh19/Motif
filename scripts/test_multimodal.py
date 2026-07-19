#!/usr/bin/env python3
"""
scripts/test_multimodal.py — End-to-end smoke test for all 5 ingestion parsers.

Verifies that each parser can:
  1. Parse a real sample file
  2. Return at least one non-empty ParsedPage
  3. Chunker can process those pages into Chunks

This does NOT use the full RAG pipeline. It tests ONLY the ingestion layer
(parser → chunker) to validate multimodal parsing without needing models
or a running Qdrant instance.

Usage:
    python scripts/test_multimodal.py                  # test all parsers
    python scripts/test_multimodal.py --parser pdf     # test one parser
    python scripts/test_multimodal.py --verbose        # show chunk counts

Requires sample files in tests/fixtures/:
    tests/fixtures/sample.pdf
    tests/fixtures/sample.docx
    tests/fixtures/sample.md
    tests/fixtures/sample.png
    tests/fixtures/sample.mp3

Run `python scripts/test_multimodal.py --create-fixtures` to generate
minimal synthetic fixtures for the test (no real files needed).
"""
from __future__ import annotations

import argparse
import sys
import time
from pathlib import Path

PROJECT_ROOT = Path(__file__).parent.parent.resolve()
FIXTURES_DIR = PROJECT_ROOT / "tests" / "fixtures"
sys.path.insert(0, str(PROJECT_ROOT))


# ── Fixture creation ──────────────────────────────────────────────────────────

def _create_fixtures() -> None:
    """Create minimal synthetic fixture files for the smoke test."""
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)

    # Markdown (trivial)
    md_file = FIXTURES_DIR / "sample.md"
    if not md_file.exists():
        md_file.write_text(
            "# Sample Document\n\n"
            "This is a test document for the multimodal ingestion smoke test.\n\n"
            "## Section 2\n\nSecond section text.\n",
            encoding="utf-8",
        )
        print(f"  created: {md_file.name}")

    # Plain text
    txt_file = FIXTURES_DIR / "sample.txt"
    if not txt_file.exists():
        txt_file.write_text(
            "This is a plain text fixture for testing the TXT parser.\n",
            encoding="utf-8",
        )
        print(f"  created: {txt_file.name}")

    # DOCX (via python-docx)
    try:
        from docx import Document
        docx_file = FIXTURES_DIR / "sample.docx"
        if not docx_file.exists():
            doc = Document()
            doc.add_heading("Sample DOCX", level=1)
            doc.add_paragraph("This is a synthetic DOCX fixture for the smoke test.")
            doc.add_paragraph("Second paragraph with more content.")
            doc.save(str(docx_file))
            print(f"  created: {docx_file.name}")
    except ImportError:
        print("  [skip] python-docx not installed — cannot create sample.docx")

    # PNG (minimal 1×1 white pixel via stdlib)
    try:
        import struct, zlib
        png_file = FIXTURES_DIR / "sample.png"
        if not png_file.exists():
            # Minimal valid 8x8 white PNG
            def _make_png() -> bytes:
                header = b"\x89PNG\r\n\x1a\n"
                ihdr = struct.pack(">IIBBBBB", 8, 8, 8, 2, 0, 0, 0)
                ihdr_chunk = b"IHDR" + ihdr
                ihdr_crc = struct.pack(">I", zlib.crc32(ihdr_chunk) & 0xffffffff)

                raw = b"".join(b"\x00" + b"\xff\xff\xff" * 8 for _ in range(8))
                compressed = zlib.compress(raw)
                idat_chunk = b"IDAT" + compressed
                idat_crc = struct.pack(">I", zlib.crc32(idat_chunk) & 0xffffffff)

                iend_chunk = b"IEND"
                iend_crc = struct.pack(">I", zlib.crc32(iend_chunk) & 0xffffffff)

                def _pack_chunk(data: bytes) -> bytes:
                    name = data[:4]
                    body = data[4:]
                    return struct.pack(">I", len(body)) + data + struct.pack(">I", zlib.crc32(data) & 0xffffffff)

                return (
                    header
                    + _pack_chunk(ihdr_chunk)
                    + _pack_chunk(idat_chunk)
                    + _pack_chunk(iend_chunk)
                )
            png_file.write_bytes(_make_png())
            print(f"  created: {png_file.name}")
    except Exception as e:
        print(f"  [skip] PNG fixture creation failed: {e}")

    print("Fixtures ready in:", FIXTURES_DIR)


# ── Parser test runners ────────────────────────────────────────────────────────

def _test_parser(
    parser_name: str,
    fixture_name: str,
    verbose: bool,
) -> tuple[bool, str, float]:
    """
    Test a single parser.

    Returns:
        (passed, message, elapsed_ms)
    """
    from rag.config import RAGConfig
    from rag.ingestion.chunker import SentenceChunker, ChunkerConfig

    fixture = FIXTURES_DIR / fixture_name
    if not fixture.exists():
        return False, f"fixture not found: {fixture}", 0.0

    config = RAGConfig()

    try:
        from rag.ingestion.parsers.base import get_parser
        parser = get_parser(fixture, config)
    except Exception as exc:
        return False, f"get_parser() failed: {exc}", 0.0

    t0 = time.monotonic()
    try:
        pages = parser.parse(fixture)
    except Exception as exc:
        elapsed = (time.monotonic() - t0) * 1000
        return False, f"parse() raised: {exc}", elapsed

    elapsed = (time.monotonic() - t0) * 1000

    if not pages:
        return False, "parse() returned 0 pages", elapsed

    total_chars = sum(len(p.text) for p in pages)
    if total_chars < 5:
        return False, f"parse() returned {len(pages)} pages but total text is only {total_chars} chars", elapsed

    # Chunker test
    chunker = SentenceChunker(ChunkerConfig(target_tokens=256, overlap_tokens=32))
    chunks = chunker.chunk_pages(
        pages,
        source=str(fixture),
        filename=fixture.name,
        source_type=parser_name,
    )

    msg = (
        f"{len(pages)} pages, {len(chunks)} chunks, {total_chars} chars"
        if verbose else f"{len(pages)} pages → {len(chunks)} chunks"
    )
    return True, msg, elapsed


# ── Main ─────────────────────────────────────────────────────────────────────

_PARSER_FIXTURES: list[tuple[str, str]] = [
    ("pdf",    "sample.pdf"),
    ("docx",   "sample.docx"),
    ("md",     "sample.md"),
    ("image",  "sample.png"),
    ("audio",  "sample.mp3"),
]


def main() -> None:
    arg_parser = argparse.ArgumentParser(
        description="End-to-end smoke test for all multimodal parsers."
    )
    arg_parser.add_argument("--parser", choices=["pdf", "docx", "md", "image", "audio"],
                            help="Test only this parser")
    arg_parser.add_argument("--verbose", "-v", action="store_true", help="Show extra details")
    arg_parser.add_argument("--create-fixtures", action="store_true",
                            help="Generate synthetic fixture files and exit")
    args = arg_parser.parse_args()

    if args.create_fixtures:
        print("Creating synthetic fixtures...")
        _create_fixtures()
        sys.exit(0)

    to_test = (
        [(name, fix) for name, fix in _PARSER_FIXTURES if name == args.parser]
        if args.parser else _PARSER_FIXTURES
    )

    passed = 0
    failed = 0
    skipped = 0

    print(f"\n{'Parser':<10}  {'Status':<8}  {'Time':>8}  Info")
    print("-" * 65)

    for name, fixture in to_test:
        ok, msg, elapsed = _test_parser(name, fixture, args.verbose)
        if "fixture not found" in msg:
            status = "SKIP"
            skipped += 1
        elif ok:
            status = "PASS"
            passed += 1
        else:
            status = "FAIL"
            failed += 1

        color = {
            "PASS": "\033[32m",
            "FAIL": "\033[31m",
            "SKIP": "\033[33m",
        }.get(status, "")
        reset = "\033[0m"

        print(f"{name:<10}  {color}{status}{reset:<8}  {elapsed:>6.0f} ms  {msg}")

    print("-" * 65)
    total = len(to_test)
    print(f"Total: {total}  Passed: {passed}  Failed: {failed}  Skipped: {skipped}\n")

    if failed:
        sys.exit(1)


if __name__ == "__main__":
    main()
