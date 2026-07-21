import math
import os
import struct
import wave

import docx
import fitz  # PyMuPDF


def create_fixtures():
    out_dir = "tests/fixtures"
    os.makedirs(out_dir, exist_ok=True)
    
    # 1. WAV audio
    wav_path = os.path.join(out_dir, "sample.wav")
    with wave.open(wav_path, 'w') as w:
        w.setnchannels(1)
        w.setsampwidth(2)
        w.setframerate(16000)
        # generate 1 second of 440hz tone
        for i in range(16000):
            value = int(10000.0 * math.cos(440.0 * math.pi * float(i) / 16000.0))
            data = struct.pack('<h', value)
            w.writeframesraw(data)
    print(f"Created: {wav_path}")

    # 2. Complex PDF
    pdf_path = os.path.join(out_dir, "sample.pdf")
    doc = fitz.open()
    page = doc.new_page()  # type: ignore
    page.insert_text((50, 50), "This is a complex PDF.")
    page.insert_text((50, 100), "It contains multiple sections and simulated tabular data.")
    page.insert_text((50, 150), "Section 1\nRow 1: 100\nRow 2: 200")
    page.insert_text((50, 200), "Summary: Key findings indicate a 50% increase.")
    doc.save(pdf_path)
    print(f"Created: {pdf_path}")
    
    # 3. DOCX
    docx_path = os.path.join(out_dir, "sample.docx")
    doc_x = docx.Document()
    doc_x.add_heading("Sample DOCX", level=1)
    doc_x.add_paragraph("This is a synthetic DOCX fixture for the smoke test.")
    doc_x.add_heading("Main sections", level=2)
    doc_x.add_paragraph("Second paragraph with more content about this document.")
    doc_x.save(docx_path)
    print(f"Created: {docx_path}")
    
    # 4. Markdown
    md_path = os.path.join(out_dir, "sample.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("# Sample Document\n\nThis document covers various topics.\n\n## Section 2\n\nThe key points in section 2 are A, B, and C.\n")
    print(f"Created: {md_path}")
    
    # 5. Image (PNG)
    png_path = os.path.join(out_dir, "sample.png")
    # minimal valid png
    import zlib
    def _make_png() -> bytes:
        header = b"\x89PNG\r\n\x1a\n"
        ihdr = struct.pack(">IIBBBBB", 8, 8, 8, 2, 0, 0, 0)
        ihdr_chunk = b"IHDR" + ihdr
        ihdr_crc = struct.pack(">I", zlib.crc32(ihdr_chunk) & 0xffffffff)
        raw = b"".join(b"\x00" + b"\xff\xff\xff" * 8 for _ in range(8))
        compressed = zlib.compress(raw)
        idat_chunk = b"IDAT" + compressed
        idat_crc = struct.pack(">I", zlib.crc32(idat_chunk) & 0xffffffff)
        iend_chunk = b"IEND"
        iend_crc = struct.pack(">I", zlib.crc32(iend_chunk) & 0xffffffff)
        def _pack_chunk(data: bytes) -> bytes:
            return struct.pack(">I", len(data) - 4) + data + struct.pack(">I", zlib.crc32(data) & 0xffffffff)
        return header + _pack_chunk(ihdr_chunk) + _pack_chunk(idat_chunk) + _pack_chunk(iend_chunk)
    with open(png_path, "wb") as f:
        f.write(_make_png())
    print(f"Created: {png_path}")

if __name__ == "__main__":
    create_fixtures()
