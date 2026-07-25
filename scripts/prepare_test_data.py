import os
import urllib.request
import json
from pathlib import Path
import docx

DATA_DIR = Path("tests/evaluation/definitive_test_data")
os.makedirs(str(DATA_DIR), exist_ok=True)

def download_file(url: str, filename: str):
    path = DATA_DIR / filename
    if not path.exists():
        print(f"Downloading {filename}...")
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
            out_file.write(response.read())
    else:
        print(f"{filename} already exists.")
    return path

def create_docx(filename: str):
    path = DATA_DIR / filename
    if not path.exists():
        print(f"Generating {filename}...")
        doc = docx.Document()
        doc.add_heading('Project Alpha Operations Manual', 0)
        doc.add_heading('Section 1: Budget and Planning', level=1)
        doc.add_paragraph('Project Alpha is the flagship initiative for Q3. The total allocated budget for this project is strictly capped at $4.2 million USD. Any expenditures exceeding this must be approved by the board.')
        
        doc.add_heading('Section 2: Timeline', level=1)
        doc.add_paragraph('The kickoff is scheduled for October 15th. We expect the beta launch to occur by January 10th of the following year.')
        doc.add_paragraph('The team lead for the beta launch is Robert Oppenheimer, based in the New York office.')
        doc.save(path)
    return path

def create_audio(filename: str):
    path = DATA_DIR / filename
    if not path.exists():
        print(f"Generating {filename}...")
        from gtts import gTTS
        tts = gTTS('This is a spoken audio test for the Motif RAG system. The system should correctly process this sentence and answer questions about it. The secret code word is Antigravity.')
        tts.save(str(path))
    return path

def create_synthetic_files():
    import csv
    import numpy as np
    import cv2
    
    # MD
    md_path = DATA_DIR / "project_zeta.md"
    if not md_path.exists():
        with open(md_path, "w", encoding="utf-8") as f:
            f.write("# Project Zeta\n\n## Dependencies\nWe rely on ZetaLib v2.0 for all our core processing. It is strictly required.")
            
    # HTML
    html_path = DATA_DIR / "corporate_policy.html"
    if not html_path.exists():
        with open(html_path, "w", encoding="utf-8") as f:
            f.write("<html><head><title>Acme Corp Policy HTML Document</title></head><body><h1>Financial Guidelines</h1><p>The Chief Financial Officer (CFO) is Jane Doe. The maximum reimbursable daily expense limit is exactly $50.</p></body></html>")
            
    # CSV
    csv_path = DATA_DIR / "synthetic_sales.csv"
    if not csv_path.exists():
        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["Date", "Item", "Amount", "Status"])
            writer.writerow(["2024-01-01", "Widget A", "150", "Completed"])
            writer.writerow(["2024-01-02", "Widget B", "200", "Pending"])
            writer.writerow(["2024-01-03", "Widget C", "99", "Completed"])
            
    # Image (Text Image for OCR)
    img_path = DATA_DIR / "test_text.png"
    if not img_path.exists():
        img = np.ones((200, 1000, 3), dtype=np.uint8) * 255
        cv2.putText(img, "The word written on this image is CONFIDENTIAL", (10, 100), cv2.FONT_HERSHEY_SIMPLEX, 1.2, (0, 0, 0), 3)
        cv2.imwrite(str(img_path), img)

def main():
    download_file("https://arxiv.org/pdf/1810.04805.pdf", "bert_paper.pdf")
    
    create_docx("project_alpha.docx")
    create_audio("audio_test.mp3")
    create_synthetic_files()

    ground_truth = [
        # PDF (BERT Paper)
        {
            "query": "In the BERT paper, what does BERT stand for?",
            "expected_answer": "Bidirectional Encoder Representations from Transformers",
            "format": "pdf"
        },
        {
            "query": "According to the BERT paper abstract, what is the size of the BERT_BASE model in terms of parameters?",
            "expected_answer": "110M parameters",
            "format": "pdf"
        },
        {
            "query": "What two pre-training tasks are used for BERT?",
            "expected_answer": "Masked language model (MLM) and next sentence prediction (NSP)",
            "format": "pdf"
        },
        # MD (Project Zeta)
        {
            "query": "What version of ZetaLib does Project Zeta rely on?",
            "expected_answer": "ZetaLib v2.0",
            "format": "md"
        },
        {
            "query": "Is ZetaLib strictly required for Project Zeta?",
            "expected_answer": "Yes",
            "format": "md"
        },
        {
            "query": "What is the name of the project?",
            "expected_answer": "Project Zeta",
            "format": "md"
        },
        # CSV (Synthetic Sales)
        {
            "query": "What is the amount for Widget B?",
            "expected_answer": "200",
            "format": "csv"
        },
        {
            "query": "What is the status of Widget C?",
            "expected_answer": "Completed",
            "format": "csv"
        },
        {
            "query": "Which item was sold on 2024-01-01?",
            "expected_answer": "Widget A",
            "format": "csv"
        },
        # HTML (Corporate Policy)
        {
            "query": "Who is the Chief Financial Officer?",
            "expected_answer": "Jane Doe",
            "format": "html"
        },
        {
            "query": "What is the maximum reimbursable daily expense limit?",
            "expected_answer": "$50",
            "format": "html"
        },
        {
            "query": "What is the title of the HTML policy document?",
            "expected_answer": "Acme Corp Policy",
            "format": "html"
        },
        # Image (Test Text)
        {
            "query": "What word is written on the image?",
            "expected_answer": "CONFIDENTIAL",
            "format": "image"
        },
        {
            "query": "Is the word CONFIDENTIAL written in all caps?",
            "expected_answer": "Yes",
            "format": "image"
        },
        {
            "query": "Does the image contain any financial data?",
            "expected_answer": "I cannot find an answer to this in the available documents.",
            "format": "image"
        },
        # Audio (gTTS Audio Test)
        {
            "query": "In the audio test, what is the secret code word?",
            "expected_answer": "Antigravity",
            "format": "audio"
        },
        {
            "query": "What is the spoken audio testing?",
            "expected_answer": "The Motif RAG system.",
            "format": "audio"
        },
        {
            "query": "Does the speaker expect the system to process the sentence correctly?",
            "expected_answer": "Yes.",
            "format": "audio"
        },
        # DOCX (Project Alpha)
        {
            "query": "What is the total allocated budget for Project Alpha?",
            "expected_answer": "$4.2 million USD",
            "format": "docx"
        },
        {
            "query": "When is the beta launch for Project Alpha scheduled to occur?",
            "expected_answer": "January 10th of the following year.",
            "format": "docx"
        },
        {
            "query": "Who is the team lead for the beta launch?",
            "expected_answer": "Robert Oppenheimer",
            "format": "docx"
        }
    ]

    with open(DATA_DIR / "ground_truth.json", "w") as f:
        json.dump(ground_truth, f, indent=4)
        
    print(f"Test data and {len(ground_truth)} ground truth questions prepared successfully.")

if __name__ == "__main__":
    main()
