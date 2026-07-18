import pytest
from pathlib import Path
from rag.ingestion.parsers.docx import DOCXParser
from rag.ingestion.parsers.base import ParsedPage

# If python-docx is not installed, we should skip tests or fail gracefully
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

@pytest.fixture
def docx_with_headings_and_table(tmp_path):
    if not HAS_DOCX:
        pytest.skip("python-docx not installed")
        
    doc = docx.Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the introduction.")
    doc.add_heading("Methods", level=2)
    doc.add_paragraph("These are the methods used.")
    
    table = doc.add_table(rows=2, cols=3)
    row0 = table.rows[0].cells
    row0[0].text = "A"
    row0[1].text = "B"
    row0[2].text = "C"
    row1 = table.rows[1].cells
    row1[0].text = "1"
    row1[1].text = "2"
    row1[2].text = "3"
    
    p = tmp_path / "test.docx"
    doc.save(str(p))
    return p

def test_docx_parser_extracts_headings_as_sections(docx_with_headings_and_table):
    parser = DOCXParser()
    pages = parser.parse(docx_with_headings_and_table)
    
    sections = [p.section for p in pages if p.section]
    assert "Introduction" in sections
    assert "Methods" in sections

def test_docx_parser_converts_table_to_markdown(docx_with_headings_and_table):
    parser = DOCXParser()
    pages = parser.parse(docx_with_headings_and_table)
    
    assert any("|" in p.text for p in pages)
    assert any(p.has_table for p in pages)

def test_docx_parser_extension():
    assert DOCXParser.can_parse(Path("doc.docx")) is True
    assert DOCXParser.can_parse(Path("doc.pdf")) is False
