import os
import wave
from PIL import Image, ImageDraw
import docx
import fitz  # pymupdf

data_dir = r"C:\Users\omen\OneDrive\Desktop\Motif\tests\evaluation\test_data"
os.makedirs(data_dir, exist_ok=True)

# 1. Create PDF
pdf_path = os.path.join(data_dir, "test.pdf")
doc = fitz.open()
page = doc.new_page()
page.insert_text((50, 50), "This is a test PDF document for Motif RAG.")
page.insert_text((50, 100), "It contains multiple lines to test chunking.")
doc.save(pdf_path)
doc.close()
print("Created test.pdf")

# 2. Create DOCX
docx_path = os.path.join(data_dir, "test.docx")
doc_word = docx.Document()
doc_word.add_paragraph("This is a test DOCX document for Motif RAG.")
doc_word.add_paragraph("Testing word document parsing capabilities.")
doc_word.save(docx_path)
print("Created test.docx")

# 3. Create Image (PNG)
img_path = os.path.join(data_dir, "test.png")
img = Image.new("RGB", (300, 100), color="white")
draw = ImageDraw.Draw(img)
draw.text((10, 10), "Test Image with OCR Text", fill="black")
img.save(img_path)
print("Created test.png")

# 4. Create Audio (WAV)
wav_path = os.path.join(data_dir, "test.wav")
# Create a dummy silent wav file
with wave.open(wav_path, 'wb') as wav:
    wav.setnchannels(1)
    wav.setsampwidth(2)
    wav.setframerate(44100)
    # write 1 second of silence
    wav.writeframes(b'\x00' * (44100 * 2))
print("Created test.wav")
